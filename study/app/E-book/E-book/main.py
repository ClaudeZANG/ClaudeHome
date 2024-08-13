import os
import docx
import requests

# 读取Word文档的文本内容
def read_docx(file_path):
    doc = docx.Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

# 使用DeepL翻译文本
def translate_text(text, target_lang='EN'):
    # 使用你的DeepL API密钥替换YOUR_DEEPL_API_KEY
    deepl_api_key = 'YOUR_DEEPL_API_KEY'
    url = 'https://api-free.deepl.com/v2/translate'
    params = {
        'auth_key': deepl_api_key,
        'text': text,
        'target_lang': target_lang
    }
    response = requests.post(url, data=params)
    if response.status_code == 200:
        return response.json()['translations'][0]['text']
    else:
        return None

# 创建新的Word文档并写入翻译好的文本内容
def create_translated_docx(original_file, translated_text, output_dir):
    original_filename = os.path.basename(original_file)
    translated_filename = os.path.splitext(original_filename)[0] + 'e.docx'
    translated_file = os.path.join(output_dir, translated_filename)
    document = docx.Document()
    document.add_paragraph(translated_text, style='BodyText')
    document.save(translated_file)

# 主函数
def main():
    input_dir = 'input_directory'  # 输入文件夹路径
    output_dir = 'output_directory'  # 输出文件夹路径

    # 检查输出文件夹是否存在，如果不存在则创建
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    # 遍历输入文件夹中的所有Word文档
    for root, dirs, files in os.walk(input_dir):
        for file in files:
            if file.endswith('.docx'):
                file_path = os.path.join(root, file)
                # 读取原始文档内容
                original_text = read_docx(file_path)
                # 翻译文本
                translated_text = translate_text(original_text)
                # 创建并保存翻译好的文档
                create_translated_docx(file_path, translated_text, output_dir)

if __name__ == "__main__":
    main()

